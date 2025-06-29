"""
python-docx操作word文档
"""

import docx
from docx.shared import Inches  # 定义插入图片的英寸

# 获取文档对象
document = docx.Document()

# 新增段落
p2 = document.add_paragraph('这是个段落')

# 在p2段落前插入一个段落
p1 = p2.insert_paragraph_before('第一个段落')

# 新增标题
document.add_heading('这是标题', level=1)

# 新增分页符
document.add_page_break()

# 新增表格
table = document.add_table(rows=6, cols=6)
cell = table.cell(0, 2)
cell.text = '第一行第三列'
row = table.rows[1]
row.cells[0].text = '第二行第一列'
row.cells[1].text = '第二行第二列'

# 添加图片   Inches(1)：1英寸
document.add_picture('liuyifei.jpg', width=Inches(1))
document.save('new.docx')
