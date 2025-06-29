"""
利用pandas python-docx操作word，pandas操作excel,并生成工作周报
"""

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import drawData

"""1.操作excel,将数据报表保存为图片"""
result_tuple = drawData.draw_bar('people.xlsx', 'Students', 'name', 'score')
plt = result_tuple[0]
students = result_tuple[1]
print(students)
image_name = 'data.jpg'
# 保存图片
plt.savefig(image_name)

"""2.操作word"""
document = Document()
# 添加标题
document.add_heading('数据分析报告', level=0)

# 返回不是标题行的第一行中的name列值
first_student_name = students.iloc[0]['name']
first_student_score = students.iloc[0]['score']
print(first_student_name, first_student_score) # lisi 90

# 向文档中添加段落
p = document.add_paragraph('分数排在第一个的学生是')
# 向段落末尾添加内容，并且加粗
p.add_run(str(first_student_name)).bold = True
p.add_run(',分数为')
p.add_run(str(first_student_score)).bold = True

document.add_paragraph(f'总共有{len(students.name)}名学生参加了考试，学生考试的总体情况：')

# 向文档中添加表格
table = document.add_table(rows=len(students.name) + 1, cols=2)
table.style = 'LightShading-Accent1'

# 设置表格标题
table.cell(0, 0).text = '学生姓名'
table.cell(0, 1).text = '学生分数'

# enumerate: 用于迭代过程中同时获取索引和值
for i, (index, row) in enumerate(students.iterrows()):
    table.cell(i+1, 0).text = str(row['name'])
    table.cell(i+1, 1).text = str(row['score'])

document.add_paragraph('')
document.add_picture(image_name, width=Inches(5))
document.save('Students.docx')
print('Done!!!')
